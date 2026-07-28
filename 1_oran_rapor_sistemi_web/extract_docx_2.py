import sys
import json
from docx import Document

def main():
    doc_path = "d:\\staj\\faaliyetraportaslakalma\\haftalık rapor şablon.docx"
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        return

    content = []
    
    # Extract core properties
    core_props = doc.core_properties
    content.append(f"Title: {core_props.title}")
    
    # Extract paragraphs
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            content.append(f"Para {i}: {text}")
            
    # Extract tables
    for i, table in enumerate(doc.tables):
        content.append(f"--- Table {i} ---")
        for r, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            content.append(f"Row {r}: {' | '.join(row_data)}")
            
    print(json.dumps(content, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    main()
